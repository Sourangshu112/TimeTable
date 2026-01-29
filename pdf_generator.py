import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os

# --- CONFIGURATION ---
INPUT_FILE = 'timetable.csv'
TEMPLATE_FILE = 'template_1.docx' # Must have 10 columns (Day + S1..S9)
OUTPUT_DIR_CLASS = os.path.join('Merged_Timetables', 'Class')
OUTPUT_DIR_TEACHER = os.path.join('Merged_Timetables', 'Teacher')

os.makedirs(OUTPUT_DIR_CLASS, exist_ok=True)
os.makedirs(OUTPUT_DIR_TEACHER, exist_ok=True)

def replace_placeholder_everywhere(doc, placeholder, replacement):
    """Replaces text in Body, Tables, Headers, and Text Boxes."""
    def replace_in_paragraph(paragraph):
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)

    # 1. Body & Tables
    for p in doc.paragraphs: replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: replace_in_paragraph(p)
    
    # 2. Headers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs: replace_in_paragraph(p)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs: replace_in_paragraph(p)

    # 3. Text Boxes
    for txbx in doc.element.body.findall('.//' + qn('w:txbxContent')):
        for element in txbx:
            if element.tag == qn('w:p'):
                from docx.text.paragraph import Paragraph
                replace_in_paragraph(Paragraph(element, doc))
            elif element.tag == qn('w:tbl'):
                 from docx.table import Table
                 t = Table(element, doc)
                 for row in t.rows:
                     for cell in row.cells:
                         for p in cell.paragraphs: replace_in_paragraph(p)

def merge_cells_logic(table_row, data_row):
    """Merges consecutive cells with identical content."""
    # 1. Fill Day
    cell_day = table_row.cells[0]
    cell_day.text = str(data_row['Day'])
    if cell_day.paragraphs:
        cell_day.paragraphs[0].runs[0].font.bold = True

    # 2. Iterate Slots S1-S9
    slot_cols = [f'S{i}' for i in range(1, 10)]
    start_col_idx = 1 
    previous_content = str(data_row.get('S1', ''))
    
    for i in range(1, len(slot_cols)):
        current_slot = slot_cols[i]
        current_col_idx = i + 1 
        current_content = str(data_row.get(current_slot, ''))

        if current_content != previous_content:
            first_cell = table_row.cells[start_col_idx]
            last_cell = table_row.cells[current_col_idx - 1]
            if first_cell != last_cell:
                first_cell.merge(last_cell)
            
            first_cell.text = previous_content
            for paragraph in first_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            start_col_idx = current_col_idx
            previous_content = current_content

    # Final block
    first_cell = table_row.cells[start_col_idx]
    last_cell = table_row.cells[len(slot_cols)]
    if first_cell != last_cell:
        first_cell.merge(last_cell)
    first_cell.text = previous_content
    for paragraph in first_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def generate_classwise(df):
    print("Generating Class-wise Timetables...")
    grouped = df.groupby('Batch Name')
    
    for batch_name, group in grouped:
        doc = Document(TEMPLATE_FILE)
        target_table = next((t for t in doc.tables if len(t.columns) >= 10), None)
        
        if not target_table: continue

        replace_placeholder_everywhere(doc, "{{ title }}", f"Timetable: {batch_name}")

        day_order = {'Mon': 0, 'Tue': 1, 'Wed': 2, 'Thu': 3, 'Fri': 4, 'Sat': 5}
        group['day_num'] = group['Day'].map(day_order)
        group = group.sort_values('day_num')

        for _, row in group.iterrows():
            new_row = target_table.add_row()
            # Clean row data for classwise view
            clean_row = {'Day': row['Day']}
            for i in range(1, 10):
                slot = f'S{i}'
                val = str(row[slot])
                if not val or val == 'nan' or 'lunch' in val.lower():
                    clean_row[slot] = ""
                else:
                    # Remove teacher code for cleaner Class View (optional)
                    # match = re.search(r'(.*)\s+\(.*\)$', val)
                    # clean_row[slot] = match.group(1).strip() if match else val
                    clean_row[slot] = val # Keep full text
            
            merge_cells_logic(new_row, clean_row)
            
        safe_name = re.sub(r'[\\/*?:"<>|]', "", str(batch_name))
        try:
            doc.save(os.path.join(OUTPUT_DIR_CLASS, f"{safe_name}.docx"))
            print(f"Saved Class: {safe_name}.docx")
        except PermissionError:
            print(f"SKIPPED: {safe_name}.docx is open. Close it!")

def generate_teacherwise(df):
    print("Generating Teacher-wise Timetables...")
    
    # 1. Parse Records
    records = []
    slot_cols = [f'S{i}' for i in range(1, 10)]
    
    for _, row in df.iterrows():
        batch = row['Batch Name']
        day = row['Day']
        for slot in slot_cols:
            val = str(row[slot])
            if not val or val == 'nan' or 'lunch' in val.lower(): continue
            
            # Extract Subject & Teacher
            match = re.search(r'(.*)\s+\(([^)]+)\)$', val)
            if match:
                subject = match.group(1).strip()
                teachers_str = match.group(2).strip()
                # Split multiple teachers "T1, T2"
                for teacher in teachers_str.split(','):
                    records.append({
                        'Teacher': teacher.strip(),
                        'Day': day,
                        'Slot': slot,
                        'Subject': subject,
                        'Batch': batch
                    })

    if not records: return
    t_df = pd.DataFrame(records)
    unique_teachers = sorted(t_df['Teacher'].unique())
    days_order = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri']

    for teacher in unique_teachers:
        doc = Document(TEMPLATE_FILE)
        target_table = next((t for t in doc.tables if len(t.columns) >= 10), None)
        if not target_table: continue

        replace_placeholder_everywhere(doc, "{{ title }}", f"Timetable: {teacher}")

        teacher_data = t_df[t_df['Teacher'] == teacher]
        
        for day in days_order:
            row_data = {'Day': day}
            day_records = teacher_data[teacher_data['Day'] == day]
            
            for slot in slot_cols:
                # Find all batches for this teacher in this slot
                slot_classes = day_records[day_records['Slot'] == slot]
                
                if slot_classes.empty:
                    row_data[slot] = ""
                    continue
                
                # --- MERGING LOGIC ---
                # 1. Group by Subject
                subject_map = {} # { "Math": ["Batch A", "Batch B"] }
                
                for _, r in slot_classes.iterrows():
                    subj = r['Subject']
                    bat = r['Batch']
                    if subj not in subject_map:
                        subject_map[subj] = []
                    if bat not in subject_map[subj]:
                        subject_map[subj].append(bat)
                
                # 2. Format the Text
                final_texts = []
                for subj, batches in subject_map.items():
                    sorted_batches = sorted(batches)
                    batch_str = ", ".join(sorted_batches)
                    
                    # Merge Logic: If [T] (Theory), combine batches
                    if "[T]" in subj:
                        result = batch_str.split(',')[0].strip().rsplit('-', 1)[0]
                        entry = f"{subj}\n({result})"
                    else:
                        # For Labs [L], keep separate or combine? 
                        # Usually Labs are separate groups, but if you want them combined too, use same logic.
                        # For now, let's combine Labs too if it's the exact same subject title.
                        entry = f"{subj}\n({batch_str})"
                    
                    final_texts.append(entry)
                
                row_data[slot] = "\n\n".join(final_texts)
            
            new_row = target_table.add_row()
            merge_cells_logic(new_row, row_data)

        safe_name = re.sub(r'[\\/*?:"<>|]', "", str(teacher))
        try:
            doc.save(os.path.join(OUTPUT_DIR_TEACHER, f"{safe_name}.docx"))
            print(f"Saved Teacher: {safe_name}.docx")
        except PermissionError:
             print(f"SKIPPED: {safe_name}.docx is open. Close it!")

if __name__ == "__main__":
    if os.path.exists(INPUT_FILE) and os.path.exists(TEMPLATE_FILE):
        try:
            df = pd.read_csv(INPUT_FILE)
        except:
            df = pd.read_csv(INPUT_FILE, encoding='latin1')
        generate_classwise(df)
        generate_teacherwise(df)
        print("\nDone!")
    else:
        print("Missing files.")